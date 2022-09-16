#!/usr/bin/env python
# coding: utf-8
These files will be opened in binary mode., read binary (rb) for PdfFileREader() and write binary (wb) PdfFileWriter()
# In[1]:


get_ipython().system('pip install PyPDF2')

Calling getPage(4) will return a Page object for page 5 since page 0 is the first pageThe total number of pages in the document is stored in the numPages attribute of a PdfFileReader object
# In[13]:


import PyPDF2

file = open('./Downloads/6.pdf', 'rb')
readpdf = PyPDF2.PdfFileReader(file)
totalpages = readpdf.numPages

print(totalpages)

 Before we obtain the page object, the pdf has to be decrypted by calling .decrypt('swordfish')#pageObj.rotateClockwise(180)

#The rotateClockwise() and rotateCounterClockwise() methods. The degrees to rotate is passed as an integer argumentParagraph Object : A document contains multiple paragraphs. A paragraph begins on a new line and contains multiple runs. The Document object contains a list of Paragraph objects for the paragraphs in the document. (A new paragraph begins whenever the user presses ENTER or RETURN while typing in a Word document.)
Run Objects : Runs are contiguous groups of characters within a paragraph with the same style
# In[14]:


#!pip install python-docx
import docx
doc = docx.Document('abc.docx')
doc.paragraphs
#By using doc.paragraphs


# In[24]:


get_ipython().system('pip install python-docx')


# In[28]:


import docx

doc = docx.Document()
doc = docx.Document('./Downloads/ch_10.doc')
doc.paragraphs


# In[ ]:


A Run object has bold, underline,italic,strike and outline variables.


# In[ ]:


Runs can be further styled using text attributes. Each attribute can be set to one of three values:
True (the attribute is always enabled, no matter what other styles are applied to the run),
False (the attribute is always disabled),
None (defaults to whatever the run’s style is set to)

True always makes the Run object bolded and False makes it always not bolded, no matter what the style’s bold setting is. None will make the Run object just use the style’s bold setting


# In[ ]:


By Calling the docx.Document() function.


# In[29]:


import docx
doc = docx.Document()

doc.add_paragraph('Hello there!')
doc.save('hellothere.docx')


# In[ ]:


integer from 0 to 4
The arguments to add_heading() are a string of the heading text and an integer from 0 to 4. The integer 0 makes the heading the Title style, which is used for the top of the document. Integers 1 to 4 are for various heading levels, with 1 being the main heading and 4 the lowest subheading

